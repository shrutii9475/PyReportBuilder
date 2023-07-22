import os
from datetime import date, datetime
import subprocess
from tkinter import *
from tkinter import Tk, Button, Label, Frame, filedialog, messagebox, ttk
from PIL import ImageTk, Image
import pandas as pd
import mysql.connector
from mysql.connector import Error
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT

class HomePage:
    def __init__(self,login_window ):
        self.login_window = login_window

        # def __init__(self):
        # GUI PART
        self.home_window = Tk()
        self.home_window.title("Home Page")
        self.home_window.geometry('900x600')
        self.home_window.resizable(0, 0)

        # bg_color = "#FFA3AC"
        self.bg_color = "lavender"
        self.home_window.configure(background=self.bg_color)

        # head_frame --------------------------------------------------------------------------------
        self.head_frame = Frame(self.home_window, bg=self.bg_color, bd=1,height=150, relief="ridge",)
        self.head_frame.pack(padx=5,pady=(10,0),ipadx=10,ipady=10,fill=X)

        # report heading  
        heading_label = Label(self.head_frame, text="Report Generator", font=("Helvetica", 30), bg=self.bg_color, fg="#00043C")
        heading_label.pack( pady=(10,0),side=TOP)

        # logout button        
        logout_button = Button(self.head_frame, text="Log Out", font=("Helvetica", 11, 'bold'), bg='red', fg='white', command=self.logout)
        logout_button.pack(padx=10,side=RIGHT)
        
        # Connect to DB button        
        # connect_button = Button(self.head_frame, text="connect to DB", font=("Helvetica",11, 'bold'), bg='red', fg='white', command=self.start_mysql)
        # connect_button.pack(padx=10,side=RIGHT)
        # self.status_label = Label(self.head_frame, text="", fg="black", bg=self.bg_color)
        # self.status_label.pack(padx=10,side=RIGHT)

        # body_frame -------------------------------------------------------------------------------
        self.body_frame = Frame(self.home_window, bg=self.bg_color, bd=1, relief="ridge", height=450)
        self.body_frame.pack(padx=5,pady=5,ipadx=10, ipady=10, fill=BOTH)
        
        # footer_frame -----------------------------------------------------------------------------
        self.footer_frame = Frame(self.home_window, bg=self.bg_color, bd=1, relief="ridge")
        self.footer_frame.pack(padx=5,pady=5,ipadx=10, ipady=10,fill=BOTH)

        #  ------------------------------------------------------------------------------------------

        self.menu()
        self.GUIofapp()

        self.host = 'localhost'
        self.username = 'root'
        self.password = 'Shruti098%'
        self.database = 'userdata'
        self.table_name = 'trainingsdb'

        # Create a mapping dictionary for column name variations
        self.column_mapping = {
            'SNo'                                      : 'ID',
            'Training Program Name'                    : 'TrainingProgramName',
            'Training Program Date'                    : 'TrainingProgramDate',
            'Name of the Officer'                      : 'Name',
            'Designation'                              : 'Designation',
            'Gender'                                   : 'Gender',
            'Email Address'                            : 'Email',
            'Mobile Number'                            : 'MobileNumber',
            'Organization'                             : 'Organization',
            'Office Address with State/UT details'     : 'OfficeAddress',
            'State'                                    : 'State',
            'Organization Sector'                      : 'OrganizationSector',
            'Organization belongs to SME/Non-SME'      : 'OrganizationCategorySMEorNonSME',
            'Officer belongs to SC/ST (Yes/No)'        : 'OfficerBelongsToSCOrST',
            'Officer belongs to PWD category (Yes/No)' : 'OfficerBelongsToPWD'
        }

        # mapping dictionary that associates month names with their respective month numbers:
        self.month_name_to_number = {
            "January": 1,
            "February": 2,
            "March": 3,
            "April": 4,
            "May": 5,
            "June": 6,
            "July": 7,
            "August": 8,
            "September": 9,
            "October": 10,
            "November": 11,
            "December": 12
        }

        self.training_name = ''
        self.training_date = ''

    def GUIofapp(self):

        # BODY frame 2
        frame2 = Frame(self.body_frame, bg=self.bg_color, bd=1, relief="ridge")
        frame2.pack(padx=(20,10), pady=20, side=LEFT, fill=BOTH, expand=True)

        # BODY frame 3
        frame3 = Frame(self.body_frame, bg=self.bg_color, bd=1, relief="ridge")
        frame3.pack(padx=(10,20), pady=20, side=LEFT, fill=BOTH, expand=True)

        # ------------------------------------------- FOOTER FRAME 
        # heading
        footer_heading = Label(self.footer_frame, text="(OR Upload an existing file into Database)", font=('Microsoft Yahei UI Light', 14, 'bold'), bg=self.bg_color, fg='#00043C')
        footer_heading.pack(padx=10,pady=10)

        #  FOOTER FRAME 1
        frame1 = Frame(self.footer_frame, bg=self.bg_color, bd=1, relief="ridge")
        frame1.pack(padx=10)

        # Add a button *Choose a File* in FOOTER FRAME 
        choose_excel_button = Button(frame1, text='Choose Excel File', font=("Helvetica", 12, 'bold'), bg='green3', fg='white', command=self.open_file)
        choose_excel_button.pack(side=LEFT)

        # Add a Label 1 in FOOTER FRAME 
        self.filename_label = Label(self.footer_frame, text="", font=('Microsoft Yahei UI Light', 9, 'bold'), wraplength=500, bg=self.bg_color) # Adjust wrap length as needed
        self.filename_label.pack(side=BOTTOM, pady=10)

        # Add a button  *Upload to DataBase*  in FOOTER FRAME 
        upload_button = Button(frame1, text='Upload', font=("Helvetica", 12, 'bold'), bg='green3', fg='white', command=self.send_data_to_database)
        upload_button.pack(side=LEFT)

        # ----------------------------------------- BODY FRAME 02
        # Add a heading 2 for frame 2 
        body_heading = Label(frame2, text="Search Trainings By Date", font=('Microsoft Yahei UI Light', 20, 'bold'), bg=self.bg_color, fg='#00043C')
        body_heading.pack(padx=10, pady=10, side='top')
                
        # Add a heading 4 in frame 2
        subheading21 = Label(frame2, text="(Choose the Month and Year)", font=('Microsoft Yahei UI Light', 9), bg=self.bg_color, fg='#00043C')
        subheading21.pack(padx=5, pady=5)

        frameinside = Frame(frame2, background=self.bg_color )
        frameinside.pack( ipadx=10, ipady=10, pady=10)

        # def generate_date():
        #     month = self.month_combo.get()
        #     year = self.year_combo.get()

        #     # Generate the complete date
        #     date_str = f"{month}/01/{year}"
        #     print("Generated Date:", date_str)

        month_label = ttk.Label(frameinside,  font=('Microsoft Yahei UI Light', 10), text="Month:",background=self.bg_color)
        month_label.pack(padx=5, side='left')
        
        self.month_combo = ttk.Combobox(frameinside, values=['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December'], width=10)
        self.month_combo.configure(font=("Helvetica", 10), background="blue")
        self.month_combo.set('month')
        self.month_combo.pack(padx=5, side='left')

        # Year Dropdown
        year_label = ttk.Label(frameinside ,font=('Microsoft Yahei UI Light', 12), text="Year:", background=self.bg_color)
        year_label.pack(padx=5,side='left')

        # Get current year
        current_year = date.today().year

        # Create a list of years from the current year up to 10 years in the future
        year_values = [str(year) for year in range(current_year, current_year - 11, -1)]

        self.year_combo = ttk.Combobox(frameinside, values=year_values, width=10)
        # self.year_combo.packside='top'()
        self.year_combo.configure(font=("Helvetica", 10), background="blue")
        self.year_combo.set('year')
        self.year_combo.pack(padx=5, side='left')

        # Serch training Button
        search_training_button = Button(frameinside, text="Search",font=("Helvetica", 10, 'bold'), bg=self.bg_color, fg="black", command=self.search_trainings)
        search_training_button.pack(padx=(20,0),  side=LEFT)

        # Add a Label 1 in FOOTER FRAME 
        self.training_name_label = Label(frame2, text="Selected Training will show here ", font=('Microsoft Yahei UI Light', 9, 'bold'), wraplength=500, bg=self.bg_color) # Adjust wrap length as needed
        self.training_name_label.pack()

        # Add a heading 4 in frame 2
        subheading22 = Label(frame2, text="(Select Training from Search results to generate Training Report OR ignore for Monthly Training Report)", font=('Microsoft Yahei UI Light', 9), wraplength=400, bg=self.bg_color, fg='#00043C')
        subheading22.pack(pady=(10,0))

        # ----------------------------------------- BODY FRAME 03

        heading3 = Label(frame3, text="Generate Report", font=('Microsoft Yahei UI Light', 20, 'bold'), bg=self.bg_color, fg='#00043C')
        heading3.pack(padx=10, pady=10, side='top')

        # Add a heading 3 in frame 3
        subheading31 = Label(frame3, text="(Select options from Dropdown)", font=('Microsoft Yahei UI Light', 9), bg=self.bg_color, fg='#00043C')
        subheading31.pack(padx=5, pady=5)

        # Create the first dropdown
        self.options_list1 = ["Select", "Gender", "State", "Organization Sector", "Organization belongs to SME/Non-SME","Officer belongs to SC/ST (Yes/No)", "Officer belongs to PWD category (Yes/No)"]
        
        self.dropdown1 = ttk.Combobox(frame3, values=self.options_list1, state="readonly")
        self.dropdown1.set("Select")
        self.dropdown1.configure(font=("Helvetica", 12), background="blue")
        self.dropdown1.pack(padx=10, pady=(10, 10))

        # Create the second dropdown
        self.options_list2 = []

        self.dropdown2 = ttk.Combobox(frame3, values=self.options_list2, state="readonly")
        self.dropdown2.set("Select")
        self.dropdown2.configure(font=("Helvetica", 12), background="blue")
        self.dropdown2.pack(padx=10, pady=(10, 10))

        self.dropdown1.bind("<<ComboboxSelected>>", self.update_dropdown2)

        # Add a button to *generate report* in frame 2
        generate_report_button = Button(frame3, text="Generate Report", font=("Helvetica", 12, 'bold'), bg="#FF9800", fg="#ffffff", command = self.fetch_data)
        generate_report_button.pack(padx=10, pady=10, side=BOTTOM)

    def update_dropdown2(self, event):

        selected_option = self.dropdown1.get()

        if selected_option == "Select":
            self.options_list2 = ["Select"]
            # break
        elif selected_option == "Gender":
            self.options_list2 = ["Male", "Female"]
            # break
        elif selected_option == "State":
            self.options_list2 = ["Andhra Pradesh", "Arunachal Pradesh", "Assam", "Bihar", "Chhattisgarh", "Goa", "Gujarat", "Haryana", "Himachal Pradesh", "Jharkhand", "Karnataka", "Kerala","Madhya Pradesh", "Maharashtra", "Manipur", "Meghalaya", "Mizoram", "Nagaland", "New Delhi", "Odisha", "Punjab", "Rajasthan", "Sikkim", "Tamil Nadu", "Telangana", "Tripura", "Uttarakhand", "Uttar Pradesh", "West Bengal"]
            # break
        elif selected_option == "Organization Sector":
            self.options_list2 = ('Central Government', 'State Government', 'Defence', 'PSU', 'Finance','Banking', 'Power', 'Energy', 'Telecom', 'Transport','Manufacturing', 'LEA', 'Academia', 'Private', 'IT', 'ITeS')
            # break
        elif selected_option == "Organization belongs to SME/Non-SME":
            self.options_list2 = ["SME", "Non-SME"]
            # break
        elif selected_option == "Officer belongs to SC/ST (Yes/No)":
            self.options_list2 = ["Yes", "No"]
            # break
        elif selected_option == "Officer belongs to PWD category (Yes/No)":
            self.options_list2 = ["Yes", "No"]
            # break
        else:
            self.options_list2 = ["Select"]

        self.dropdown2["values"] = self.options_list2
        self.dropdown2.current(0)

    def menu(self):
        def myfun():
            print("File menu working")

        def quitw():
            result = messagebox.askquestion(
                "Quit", "Are you sure you want to quit?")
            if result == "yes":
                self.frame2.destroy()

        # making a menu bar
        my_menu = Menu(self.home_window)
        my_menu.add_command(label="File", command=myfun)
        # my_menu.add_command(label="Home", command=open_second_script)
        my_menu.add_command(label="Quit", command=quitw)
        my_menu.add_command(label="Help", command=myfun)
        self.home_window.config(menu=my_menu)

    def open_file(self):

        filename = filedialog.askopenfilename(title="Select Excel file", filetypes=[("Excel files", "*.xlsx;*.xls")])

        if filename:

            # self.data = pd.read_excel(filename)
            # self.send_data_to_database(self.data)
            self.excel_data = pd.read_excel(filename)

            # df.empty: This condition checks if the DataFrame df is empty, meaning it does not contain any rows or columns. It is typically used to check if the DataFrame has any data.
            if self.excel_data.empty:
                # messagebox.showinfo(title="Failure", message="Excel file is empty!!")
                self.filename_label.config(text="Excel file is empty!, Choose another file!", fg='black')
                return

            self.filename_label.config(text="Selected file: " + filename, fg='black')
            # messagebox.showinfo(title="Success", message="Excel file opened successfully.")
            
        else:
            del self.excel_data
            self.filename_label.config(text="❌ File Loading Failed", fg="red")
            messagebox.showinfo(title="Failure", message="Excel file opening failed. Try Again!!!")

            # self.tick_mark_label.config(text="❌ File Loading Failed", fg="red")

    def send_data_to_database(self):

        # Connect to the MySQL database
        conn = mysql.connector.connect( host = self.host, user = self.username, password = self.password, database = self.database )
        cursor = conn.cursor()

        insert_query = ''' INSERT INTO trainingsdb (
            TrainingProgramName,
            TrainingProgramDate,
            Name,
            Designation,
            Gender,
            Email,
            MobileNumber,
            Organization,
            OfficeAddress,
            State,
            OrganizationSector,
            OrganizationCategorySMEorNonSME,
            OfficerBelongsToSCOrST,
            OfficerBelongsToPWD
        ) values (
            %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
        )
        '''
        success_count = 0
        for index, row in self.excel_data.iterrows():
            record_data = []
            # for value in row:
            
            # for column_name, value in row.items():

            for column_name in row.index:
                if column_name != 'SNo':  # Skip the 'SNo' column
                    value = row[column_name]        # to comment later
                    if pd.isnull(value):
                        record_data.append(None)  # Convert NaN values to None (null)
                    else:
                        record_data.append(value)
                
            # print('index = ',index)
            # print('row = ',row)
            # print('record_data = ',record_data)
            try:
                cursor.execute(insert_query, record_data)  # Execute the query
                conn.commit() 
                success_count += 1

            except Exception as e:
                print("Error inserting record:", e)

        cursor.close()
        conn.close()

        if (success_count == 0):
            messagebox.showerror("Error","Error inserting record")
        else :
            # Display a message box with the success count
            messagebox.showinfo("Query Execution Result", f"Query executed successfully for {success_count} records.")
    
    def on_treeview_select(self, event):
        selected_item = self.tree.selection()[0]
        self.training_name = self.tree.item(selected_item, 'values')[0]  # First value is the training name
        self.training_date = self.tree.item(selected_item, 'values')[1]  # Second value is the training date
        
        # Now you can use the 'training_name' variable for other purposes, like storing it or further processing.
        self.training_name_label_frame2 = "Selected Training name:" + self.training_name + "/ Date:" + self.training_date
        print("Selected Training name:" + self.training_name + "/Training Date:" + self.training_date)

        self.training_name_label.config(text= self.training_name_label_frame2, fg='black')

    def show_result_in_listview(self, total_trainings, count):
        
        # Create a new window to show the result in a listview
        result_window = Toplevel(self.home_window)
        result_window.title("Search Results")
        result_window.geometry("600x300")

        # Create a listview widget
        columns = ("Training Program Name", "Training Program Date")
        self.tree = ttk.Treeview(result_window, columns=columns, show="headings")
        
        # Add column headings
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100, anchor="center")

        # Insert total_trainings into the listview
        for item in total_trainings:
            self.tree.insert("", "end", values=(item))

        # Pack the listview and scrollbar
        self.tree.pack(fill="both", expand=True)

        # Bind the TreeviewSelect event to the on_treeview_select function
        self.tree.bind("<<TreeviewSelect>>", self.on_treeview_select)
        
    def search_trainings(self):

        # Connect to the MySQL database
        conn = mysql.connector.connect( host = self.host, user = self.username, password = self.password, database = self.database )
        cursor = conn.cursor()

        month = self.month_combo.get()
        year = self.year_combo.get()

        # month_number = home_page.month_name_to_number.get(month)
        month_number = self.month_name_to_number[month]
    
        # checks if the user has selected both a month and a year
        if (month == 'month') or (year == 'year'):
            messagebox.showinfo("Select Date", "Choose a Month and Year")
            return 

        print('Searching Trainings for the Date: ',str(month_number)+'/01/'+year)

            # SQL query to Search for the training program names from the database table.
        try: 
            query = f"SELECT TrainingProgramName,TrainingProgramDate FROM trainingsdb WHERE MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"
            cursor.execute(query)
            result = cursor.fetchall()

            # CHECK If there are no training programs found for the selected date
            if len(result) == 0:
                messagebox.showinfo('Empty Set','No training is available in the selected date!!!')
                print('Empty Set','No training is available in the selected date!!!')
            else:
                # count the number of training programs for the selected date
                count = len(result)
                print('Total number of trainings: ', count)
                messagebox.showinfo("Search Result", f"Found "+ str(count) + " Trainings for the selected date.")
                self.training_date = ''
                self.training_name = ''
                self.training_name_label.config(text= "Selected Training will show here!" , fg='black')

                self.show_result_in_listview(result, count)
                # for row in result:
                #     print(row)

        except Error as e:
            print("Error searching dates:", str(e))
            messagebox.showerror('Error searching dates:', str(e))

        conn.commit()
        cursor.close()
        conn.close()

    def fetch_data(self):

        # Obtain the user input for column name and search value
        column = self.dropdown1.get()
        value = self.dropdown2.get()
        month = self.month_combo.get()
        year = self.year_combo.get()

        if column == 'Select' or value == 'Select':
            messagebox.showinfo("Select Data", "Choose a Column and VAlue")
            return

        if month == 'month' or year == 'year':
            messagebox.showinfo("Select Date", "Choose a Month and Year")
            return

        # column_name = home_page.column_mapping.get(column)
        # month_number = home_page.month_name_to_number.get(month)

        column_name = self.column_mapping[column]
        month_number = self.month_name_to_number[month]

        # Connect to the MySQL database
        conn = mysql.connector.connect( host = self.host, user = self.username, password = self.password, database = self.database )
        cursor = conn.cursor()

        try:
            # Query to search for monthly trainings 
            if self.training_name == '': 
                
                # search_for_training_in_date = f"SELECT * FROM trainingsdb WHERE AND {column_name} = '{value}' AND MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year} ;"
                search_for_training_in_date = f"SELECT * FROM trainingsdb WHERE {column_name} = '{value}' AND MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"

                cursor.execute(search_for_training_in_date)
                result = cursor.fetchall()

                if len(result) == 0:
                    messagebox.showinfo('Empty Set', 'No Trainings available in the selected date!!!')
                    return
                else:
                    # Process and display the search results
                    result_df_table = pd.DataFrame(result)
                    print(result_df_table)
                    # for row in result:
                    #     print(row)
                
            else:
                # SQL query to search for the specified training in the given month and year
            
                # SQL query to search for rows with the given column name, value, training name, and within the given month and year
                query = f"SELECT * FROM trainingsdb WHERE TrainingProgramName = '{self.training_name}' AND {column_name} = '{value}' AND MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"

                # Execute the query
                cursor.execute(query)

                # Fetch all the rows from the result
                result = cursor.fetchall()

                if len(result) == 0:
                    messagebox.showinfo('Empty Set', 'No Records available in the selected date!!!')
                    return
                else:
                    # Process and display the search results
                    result_df_table = pd.DataFrame(result)
                    print(result_df_table)
                    # for row in result:
                    #     print(row)

        except mysql.connector.Error as e:
            print("Error executing the queries:", e)
        except Error as e:
            print("Error executing the search query:", e)
        
        cursor.close()
        conn.close()

        self.generate_doc_report(result_df_table)
    
    def generate_doc_report(self, result_df_table):

        month = self.month_combo.get()
        year = self.year_combo.get()
        column = self.dropdown1.get()
        value = self.dropdown2.get()
        name = self.training_name

        column_name = self.column_mapping[column]
        month_number = self.month_name_to_number[month]
        # column_name = home_page.column_mapping.get(column)
        # month_number = home_page.month_name_to_number.get(month)

        # Create a new Word document
        document = Document()

        # Set document page size to A4
        section = document.sections[0]
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        # Set document layout to landscape
        section.orientation = WD_ORIENT.LANDSCAPE

        
        # montly training report
        if self.training_name == '':

            str1 = 'Monthly Training Programs Report'
            # Add a # training Name title to the document
            document.add_heading(str1, 0)

            # training date
            str2 = 'Month: ' + month + ' ' + year
            document.add_paragraph(str2)

            # sql query for all data in the month or individual training/ showing the permanent details
            document.add_paragraph("The details pertaining to the candidates attending the Training Program are as follows. ")

            # Add a table with 3 rows and 4 columns
            table = document.add_table(rows=1, cols=2)
            table.style = document.styles['Table Grid']

            # Adding table headers for the first table
            header = table.rows[0].cells
            header[0].text = 'Demographic Category'
            header[1].text = 'Count'

            def without_training_name():
                # Connect to the MySQL database
                conn = mysql.connector.connect(host=self.host, user=self.username, password=self.password, database=self.database)
                cursor = conn.cursor()

                query = """
                        SELECT COUNT(*) AS total_count FROM trainingsdb WHERE MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(*) AS female_count FROM trainingsdb WHERE Gender = 'female' AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(*) AS sc_st_count FROM trainingsdb WHERE OfficerBelongsToSCOrST = 'yes' AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(DISTINCT State) AS distinct_states_count FROM trainingsdb WHERE MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                    """
                params = (month_number,year,month_number,year,month_number,year,month_number,year)
                cursor = cursor.execute(query, params, multi=True)

                # Fetch results from each query
                results = []
                for result_set in cursor:
                    results.append(result_set.fetchone()[0])
                
                # Extract individual results from the combined result set
                count_all, count_female, count_sc_st, distinct_states_count = results

                # query = f"SELECT COUNT(*) AS total_count FROM trainingsdb WHERE MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"
                # cursor.execute(query)
                # count_all =  cursor.fetchall()

                # query = f"SELECT COUNT(*) AS female_count FROM trainingsdb WHERE Gender = 'female' AND MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"
                # cursor.execute(query)
                # count_female = cursor.fetchall()
                
                # query = f"SELECT COUNT(*) AS sc_st_count FROM trainingsdb WHERE OfficerBelongsToSCOrST = 'yes' AND MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"
                # cursor.execute(query)
                # count_sc_st = cursor.fetchall()

                # query = f"SELECT COUNT(DISTINCT State) AS distinct_states_count FROM trainingsdb WHERE MONTH(TrainingProgramDate) = {month_number} AND YEAR(TrainingProgramDate) = {year};"
                # cursor.execute(query)
                # distinct_states_count = cursor.fetchall()

                # Print the results
                print("Total count:", count_all)
                print("Count of female participants:", count_female)
                print("Count of participants belonging to SC/ST:", count_sc_st)
                print("Number of distinct states:", distinct_states_count)


                # Define the data labels and corresponding values
                data_labels = ['Total count of participants', 'Count of female participants:', 'Count of participants belonging to SC/ST:', 'Number of distinct states:']
                data_values = [count_all, count_female, count_sc_st, distinct_states_count]

                # Add rows to the table dynamically using a for loop
                for label, value in zip(data_labels, data_values):
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = str(value)

                # Close the cursor and connection
                cursor.close()
                conn.close()
            without_training_name()

            # annexure1 
            document.add_heading("Annexure 1", level=1)

            document.add_paragraph()
            # Add a paragraph with the selected option and value
            document.add_paragraph("Selected Option: " + self.dropdown1.get())
            document.add_paragraph("Selected Value: " + self.dropdown2.get())

            # # counting to number of candidates
            document.add_heading("Total number of records found: " + str(len(result_df_table)))

            document.add_paragraph()  # Add an empty paragraph for spacing

            # Add the DataFrame as a table to the document
            table2 = document.add_table(rows=1, cols=15)
            table2.autofit = True
            table2.style = document.styles['Table Grid']

            # Write the column names as the table header
            for i, column_name in enumerate(self.column_mapping.keys()):
                table2.cell(0, i).text = column_name

            # Write the data from the DataFrame to the table
            for _, row in result_df_table.iterrows():
                new_row = table2.add_row()
                for i, val in enumerate(row):
                    new_row.cells[i].text = str(val)


        # individual training report
        else:

            str1 = 'Training Program Name: ' + self.training_name
            # Add a # training Name title to the document
            document.add_heading(str1, 0)
            
            # training date
            str2 = 'Date: ' + self.training_date
            document.add_paragraph(str2)

            # sql query for all data in the month or individual training/ showing the permanent details
            document.add_paragraph("The details pertaining to the candidates attending the Training Program are as follows. ")

            # Add a table with 3 rows and 4 columns
            table = document.add_table(rows=1, cols=2)
            table.style = document.styles['Table Grid']

            # Adding table headers for the first table
            header = table.rows[0].cells
            header[0].text = 'Demographic Category'
            header[1].text = 'Count'

            def with_training_name():
                # Connect to the MySQL database
                conn = mysql.connector.connect(host=self.host, user=self.username, password=self.password, database=self.database)
                cursor = conn.cursor()

                query = """
                        SELECT COUNT(*) AS total_count FROM trainingsdb WHERE TrainingProgramName = %s AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(*) AS female_count FROM trainingsdb WHERE TrainingProgramName = %s AND Gender = 'female' AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(*) AS sc_st_count FROM trainingsdb WHERE TrainingProgramName = %s AND OfficerBelongsToSCOrST = 'yes' AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                        SELECT COUNT(DISTINCT State) AS distinct_states_count FROM trainingsdb WHERE TrainingProgramName = %s AND MONTH(TrainingProgramDate) = %s AND YEAR(TrainingProgramDate) = %s;
                    """
                params = (name,month_number,year,name,month_number,year,name,month_number,year,name,month_number,year)
                cursor = cursor.execute(query, params, multi=True)

                # Fetch results from each query
                results = []
                for result_set in cursor:
                    results.append(result_set.fetchone()[0])

                # Extract individual results from the combined result set
                count_all, count_female, count_sc_st, distinct_states_count = results

                # Print the results
                print("Total count:", count_all)
                print("Count of female participants:", count_female)
                print("Count of participants belonging to SC/ST:", count_sc_st)
                print("Number of distinct states:", distinct_states_count)

                # Define the data labels and corresponding values
                data_labels = ['Total count of participants', 'Count of female participants:', 'Count of participants belonging to SC/ST:', 'Number of distinct states:']
                data_values = [count_all, count_female, count_sc_st, distinct_states_count]

                # Add rows to the table dynamically using a for loop
                for label, value in zip(data_labels, data_values):
                    row = table.add_row().cells
                    row[0].text = label
                    row[1].text = str(value)

                # Close the cursor and connection
                cursor.close()
                conn.close()
            with_training_name()

            # annexure1 
            document.add_heading("Annexure 1", level=1)

            document.add_paragraph()
            # Add a paragraph with the selected option and value
            document.add_paragraph("Selected Option: " + self.dropdown1.get())
            document.add_paragraph("Selected Value: " + self.dropdown2.get())

            # # counting to number of candidates
            document.add_heading("Total number of records found: " + str(len(result_df_table)))

            document.add_paragraph()  # Add an empty paragraph for spacing

            # Add the DataFrame as a table to the document
            table2 = document.add_table(rows=1, cols=len(result_df_table.columns))
            table2.autofit = True
            table2.style = document.styles['Table Grid']

            # Write the column names as the table header
            for i, column_name in enumerate(self.column_mapping.keys()):
                table2.cell(0, i).text = column_name

            # Write the data from the DataFrame to the table
            for _, row in result_df_table.iterrows():
                new_row = table2.add_row()
                for i, val in enumerate(row):
                    new_row.cells[i].text = str(val)


        # Custom file name
        current_time = datetime.now()
        file_name = 'Report_' + current_time.strftime('%Y-%m-%d_%H-%M-%S') + '.docx'

        # method 1 to Save the document
        document.save(file_name)
        os.startfile(file_name)


    def logout(self):
        self.home_window.destroy()
        self.login_window.deiconify()

    def run(self):
        self.home_window.mainloop()


# Example usage:
if __name__ == "__main__":
    home_page = HomePage(login_window)
    home_page = HomePage()
    home_page.run()



